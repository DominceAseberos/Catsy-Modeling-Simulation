import docx

def edit_document():
    doc = docx.Document('Project-Proposal(DES-M&S).docx')
    
    # 1. Update the assumption about reservations
    for p in doc.paragraphs:
        if "Reservation customers will still queue" in p.text:
            p.text = "• Reservation customers bypass the cashier and barista queues completely, as their order was placed via mobile pre-order prior to arrival, resulting in an immediate table seating and zero service wait time."
            break
            
    # 2. Add the Smart Forecast assumption
    p = doc.add_paragraph("• A 'Smart Forecast' pace system adjusts barista preparation times dynamically; during Peak pace (4+ orders in queue), baristas utilize batch preparation, reducing average prep time by 25%.")
    
    doc.save('Project-Proposal(DES-M&S).docx')

if __name__ == "__main__":
    edit_document()
    print("Document successfully updated.")
